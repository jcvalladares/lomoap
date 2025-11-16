from fastapi import FastAPI, HTTPException, UploadFile, File
from pydantic import BaseModel
from transformers import AutoModelForCausalLM, AutoTokenizer
import torch
import uvicorn
from contextlib import asynccontextmanager
import psutil
import os
import pathlib
import uuid
import asyncio
import time
import shutil
import tempfile
import io

# Import shared services (models, tokenizers, loading logic, GPU helpers, jobs)
from services import (
    models,
    tokenizers,
    loading_status,
    model_device_info,
    jobs,
    _log_job,
    get_gpu_info,
    detect_model_device,
    MODEL_CONFIGS,
    load_model_on_demand,
    generate_with_model,
    _fine_tune_run,
)




def get_gpu_info():
    """Get detailed GPU information"""
    if not torch.cuda.is_available():
        return {
            "cuda_available": False,
            "message": "CUDA not available - running on CPU"
        }
    
    gpu_info = {
        "cuda_available": True,
        "gpu_count": torch.cuda.device_count(),
        "current_device": torch.cuda.current_device(),
        "devices": []
    }
    
    for i in range(torch.cuda.device_count()):
        props = torch.cuda.get_device_properties(i)
        memory_allocated = torch.cuda.memory_allocated(i) / 1024**3
        memory_reserved = torch.cuda.memory_reserved(i) / 1024**3
        memory_total = props.total_memory / 1024**3
        
        gpu_info["devices"].append({
            "id": i,
            "name": props.name,
            "memory_total_gb": f"{memory_total:.2f}",
            "memory_allocated_gb": f"{memory_allocated:.2f}",
            "memory_reserved_gb": f"{memory_reserved:.2f}",
            "memory_free_gb": f"{memory_total - memory_reserved:.2f}",
            "utilization_percent": f"{(memory_reserved / memory_total) * 100:.1f}"
        })
    
    return gpu_info

def detect_model_device(model):
    """Detect which device(s) a model is loaded on"""
    if model is None:
        return {"status": "not_loaded"}
    
    device_info = {
        "devices": [],
        "primary_device": None,
        "is_distributed": False
    }
    
    # Check if model has device_map (multi-GPU)
    if hasattr(model, 'hf_device_map') and model.hf_device_map:
        device_info["is_distributed"] = True
        device_info["device_map"] = model.hf_device_map
        
        # Extract unique devices from device map
        devices = set(str(device) for device in model.hf_device_map.values())
        device_info["devices"] = list(devices)
        
        # Find primary device (usually where most layers are)
        device_counts = {}
        for device in model.hf_device_map.values():
            device_str = str(device)
            device_counts[device_str] = device_counts.get(device_str, 0) + 1
        device_info["primary_device"] = max(device_counts.keys(), key=lambda k: device_counts[k])
    
    else:
        # Single device model
        try:
            # Try to get device from first parameter
            first_param = next(model.parameters())
            device = str(first_param.device)
            device_info["devices"] = [device]
            device_info["primary_device"] = device
        except StopIteration:
            device_info["devices"] = ["unknown"]
            device_info["primary_device"] = "unknown"
    
    return device_info

# Model configurations with distilled versions
MODEL_CONFIGS = {
    "qwen": {
        "name": "Qwen/Qwen2.5-Coder-1.5B-Instruct",  # Distilled version (1.5B vs 7B)
        "description": "Qwen2.5-Coder 1.5B - Distilled, faster, smaller version",
        "config": {
            "torch_dtype": torch.float16 if torch.cuda.is_available() else "auto",
            "device_map": "auto",
            "trust_remote_code": False,
            "low_cpu_mem_usage": True
        }
    },
    "deepseek": {
        "name": "deepseek-ai/DeepSeek-Coder-V2-Lite-Base",  # Base version (smaller than Instruct)
        "description": "DeepSeek-Coder-V2 Lite Base - Compact base model",
        "config": {
            "torch_dtype": torch.float16 if torch.cuda.is_available() else "auto",
            "device_map": "auto",
            "trust_remote_code": True,
            "low_cpu_mem_usage": True
        }
    },
    "qwen_tiny": {
        "name": "Qwen/Qwen2.5-Coder-0.5B-Instruct",  # Even smaller version
        "description": "Qwen2.5-Coder 0.5B - Ultra-lightweight version",
        "config": {
            "torch_dtype": torch.float16 if torch.cuda.is_available() else "auto",
            "device_map": "auto",
            "trust_remote_code": False,
            "low_cpu_mem_usage": True
        }
    },
    "codellama": {
        "name": "codellama/CodeLlama-7b-Python-hf",  # Meta's open-source code model
        "description": "CodeLlama 7B Python - Meta's open-source Python-focused model (no chat template)",
        "config": {
            "torch_dtype": torch.float16 if torch.cuda.is_available() else "auto",
            "device_map": "auto",
            "trust_remote_code": False,
            "low_cpu_mem_usage": True
        }
    },
    "deepseek_tiny": {
        "name": "deepseek-ai/deepseek-coder-1.3b-base",  # Smallest DeepSeek Coder model
        "description": "DeepSeek-Coder 1.3B Base - Ultra-compact base model",
        "config": {
            "torch_dtype": torch.float16 if torch.cuda.is_available() else "auto",
            "device_map": "auto",
            "trust_remote_code": True,
            "low_cpu_mem_usage": True
        }
    }
}

def load_model_on_demand(model_key: str):
    """Load a specific model on first request (lazy loading)"""
    global models, tokenizers, loading_status
    
    if model_key in models and model_key in tokenizers:
        return  # Already loaded
    
    if loading_status.get(model_key) == "loading":
        raise HTTPException(status_code=503, detail=f"Model {model_key} is currently loading, please wait...")
    
    if model_key not in MODEL_CONFIGS:
        raise HTTPException(status_code=404, detail=f"Model {model_key} not available")
    
    try:
        loading_status[model_key] = "loading"
        model_info = MODEL_CONFIGS[model_key]
        model_name = model_info["name"]
        config = model_info["config"]
        
        print(f"📦 Loading {model_key.upper()} model on demand: {model_name}")
        
        # Unload other models to free memory (if needed)
        # This implements a "swap" strategy - only one model in memory at a time
        for other_key in list(models.keys()):
            if other_key != model_key:
                print(f"🔄 Unloading {other_key.upper()} model to free memory")
                del models[other_key]
                del tokenizers[other_key]
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
        
        # Load tokenizer
        tokenizers[model_key] = AutoTokenizer.from_pretrained(
            model_name,
            trust_remote_code=config.get("trust_remote_code", False)
        )
        print(f"  ✅ {model_key.upper()} tokenizer loaded")
        
        # Load model
        models[model_key] = AutoModelForCausalLM.from_pretrained(
            model_name,
            **config
        )
        
        # Detect where the model was loaded
        device_info = detect_model_device(models[model_key])
        model_device_info[model_key] = device_info
        
        print(f"  ✅ {model_key.upper()} model loaded successfully!")
        print(f"  📍 Model location: {device_info['primary_device']}")
        if device_info['is_distributed']:
            print(f"  🔄 Distributed across devices: {device_info['devices']}")
        
        # Print detailed GPU info
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            gpu_info = get_gpu_info()
            print(f"💾 GPU Memory Status:")
            for gpu in gpu_info['devices']:
                print(f"    GPU {gpu['id']} ({gpu['name']}): {gpu['memory_allocated_gb']}GB / {gpu['memory_total_gb']}GB used ({gpu['utilization_percent']}%)")
        
        loading_status[model_key] = "ready"
        
    except Exception as e:
        loading_status[model_key] = "error"
        print(f"❌ Error loading model {model_key}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to load model {model_key}: {str(e)}")

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize the API without loading models"""
    print("🚀 Starting Multi-Model Coder API with lazy loading")
    print("💡 Models will be loaded on first request to save memory")
    
    # Initialize loading status
    for model_key in MODEL_CONFIGS.keys():
        loading_status[model_key] = "not_loaded"
    
    yield
    
    # Cleanup on shutdown
    print("🔄 Shutting down and cleaning up...")
    models.clear()
    tokenizers.clear()
    loading_status.clear()
    if torch.cuda.is_available():
        torch.cuda.empty_cache()
    print("✅ Cleanup completed")

# Initialize FastAPI app with lifespan
app = FastAPI(
    title="Distilled Multi-Model Coder API", 
    description="Optimized API using distilled code generation models for faster, memory-efficient inference",
    version="2.0",
    lifespan=lifespan
)

# Pydantic models for request/response
class PromptRequest(BaseModel):
    prompt: str
    max_tokens: int = 512  # Optional parameter with default value

class TextResponse(BaseModel):
    generated_text: str
    model_used: str  # Indicate which model was used


class ScrapeFormatRequest(BaseModel):
    url: str
    output_folder: str = "downloaded_site"
    model_key: str = "deepseek_tiny"
    max_pages: int = 50
    max_chars: int = 50000
    max_new_tokens: int = 1024


class ScrapeDownloadRequest(BaseModel):
    url: str
    output_folder: str = "downloaded_site"
    max_pages: int = 50


class ScrapeProcessRequest(BaseModel):
    output_folder: str = "downloaded_site"
    model_key: str = "deepseek_tiny"
    max_chars: int = 50000
    max_new_tokens: int = 1024
    prompt_template: str = None  # If provided, should include '{content}' or will be appended before content

def generate_with_model(request: PromptRequest, model_key: str, model_name: str):
    """Helper function to generate text with a specific model"""
    try:
        # Load model on demand if not already loaded
        load_model_on_demand(model_key)
        
        model = models[model_key]
        tokenizer = tokenizers[model_key]
        
        # Check if the model supports chat templates
        has_chat_template = hasattr(tokenizer, 'chat_template') and tokenizer.chat_template is not None
        
        if has_chat_template:
            # Use chat template for models that support it (Qwen, DeepSeek)
            messages = [
                {"role": "user", "content": request.prompt}
            ]
            
            text = tokenizer.apply_chat_template(
                messages,
                tokenize=False,
                add_generation_prompt=True,
            )
        else:
            # Direct text input for models without chat templates (CodeLlama)
            text = request.prompt
        
        model_inputs = tokenizer([text], return_tensors="pt").to(model.device)
        
        # Generate text
        with torch.no_grad():
            generated_ids = model.generate(
                **model_inputs,
                max_new_tokens=request.max_tokens,
                do_sample=True,
                temperature=0.7,
                top_p=0.8,
                pad_token_id=tokenizer.eos_token_id if tokenizer.eos_token_id else tokenizer.pad_token_id
            )
        
        # Decode the generated text
        output_ids = generated_ids[0][len(model_inputs.input_ids[0]):].tolist()
        generated_text = tokenizer.decode(output_ids, skip_special_tokens=True)
        
        return TextResponse(generated_text=generated_text, model_used=model_name)
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating text with {model_name}: {str(e)}")


class ScrapeExtractRequest(BaseModel):
    output_folder: str = "downloaded_site"
    max_chars: int = 50000

@app.post("/scrape-extract")
async def scrape_extract(req: ScrapeExtractRequest):
    """Convert previously downloaded HTML files into per-page `.txt` files under `<output_folder>/pages_txt`.

    This separates the HTML downloading step from the BeautifulSoup-based extraction.
    """
    out_dir = req.output_folder
    if not os.path.isdir(out_dir):
        raise HTTPException(status_code=400, detail=f"Output folder not found: {out_dir}")

    try:
        from bs4 import BeautifulSoup
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"BeautifulSoup (bs4) is required for extraction: {e}")

    pages_txt_dir = os.path.join(out_dir, "pages_txt")
    pathlib.Path(pages_txt_dir).mkdir(parents=True, exist_ok=True)

    saved_txt_files = []
    for root, _, files in os.walk(out_dir):
        for fname in files:
            if not fname.endswith('.html'):
                continue
            html_path = os.path.join(root, fname)
            try:
                with open(html_path, 'r', encoding='utf-8') as hf:
                    soup = BeautifulSoup(hf.read(), 'html.parser')
                    for tag in soup(['script', 'style', 'noscript']):
                        tag.decompose()
                    page_text = soup.get_text(separator=' ', strip=True)

                if not page_text.strip():
                    continue

                # Truncate per-page if necessary
                if len(page_text) > req.max_chars:
                    page_text = page_text[:req.max_chars]

                rel = os.path.relpath(html_path, out_dir)
                safe_name = rel.replace(os.sep, '_')
                if safe_name.endswith('.html'):
                    safe_name = safe_name[:-5]
                txt_path = os.path.join(pages_txt_dir, f"{safe_name}.txt")
                with open(txt_path, 'w', encoding='utf-8') as tf:
                    tf.write(page_text)
                saved_txt_files.append(txt_path)
            except Exception as e:
                print(f"Failed to extract text from {html_path}: {e}")

    if not saved_txt_files:
        raise HTTPException(status_code=500, detail="No page text files were created during extraction")

    return {
        "message": "Extraction complete, per-page txt files saved",
        "output_folder": out_dir,
        "pages_txt": saved_txt_files
    }


@app.post("/scrape-download")
async def scrape_download(req: ScrapeDownloadRequest):
    """Crawl a website and save per-page HTML files under `<output_folder>`.

    Extraction (HTML -> .txt) is intentionally separated into `/scrape-extract`.
    """
    out_dir = req.output_folder
    pathlib.Path(out_dir).mkdir(parents=True, exist_ok=True)

    # If the provided URL points directly to a PDF, download the PDF instead of crawling
    try:
        import requests
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"requests is required for downloading: {e}")

    try:
        is_pdf = False
        parsed = None
        try:
            from urllib.parse import urlparse
            parsed = urlparse(req.url)
            if parsed.path.lower().endswith('.pdf'):
                is_pdf = True
        except Exception:
            pass

        # If not obvious from URL, conservatively try a HEAD request to check content-type
        if not is_pdf:
            try:
                head = requests.head(req.url, allow_redirects=True, timeout=8)
                ctype = head.headers.get('Content-Type', '')
                if 'application/pdf' in ctype.lower():
                    is_pdf = True
            except Exception:
                # ignore HEAD failures and fallback to HTML crawl
                is_pdf = False

        if is_pdf:
            # Download the PDF into pages_pdf/
            pages_pdf_dir = os.path.join(out_dir, 'pages_pdf')
            pathlib.Path(pages_pdf_dir).mkdir(parents=True, exist_ok=True)
            try:
                resp = requests.get(req.url, stream=True, timeout=20)
                resp.raise_for_status()
                # Derive filename from path
                fname = os.path.basename(parsed.path) if parsed and parsed.path else 'downloaded.pdf'
                if not fname.lower().endswith('.pdf'):
                    fname = fname + '.pdf'
                save_path = os.path.join(pages_pdf_dir, fname)
                with open(save_path, 'wb') as wf:
                    for chunk in resp.iter_content(chunk_size=8192):
                        if chunk:
                            wf.write(chunk)
                saved_html = [save_path]
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Failed to download PDF: {e}")
        else:
            # Lazy import scraping utility
            try:
                from scrapping import crawl_website
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Failed to import scrapping utilities: {e}")

            try:
                crawl_website(req.url, output_dir=out_dir, max_pages=req.max_pages)
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Crawling failed: {str(e)}")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Download error: {e}")

    # Return list of saved HTML files for visibility
    saved_html = []
    for root, _, files in os.walk(out_dir):
        for fname in files:
            if fname.endswith('.html'):
                saved_html.append(os.path.join(root, fname))

    if not saved_html:
        return {
            "message": "Crawl completed but no HTML files were detected",
            "output_folder": out_dir,
            "html_files": []
        }

    return {
        "message": "Crawl complete, HTML files saved",
        "output_folder": out_dir,
        "html_files": saved_html
    }


@app.post("/extract-pdf")
async def extract_pdf(file: UploadFile = File(...), output_folder: str = "downloaded_site"):
    """Upload a PDF file and extract text from every page.

    Saves the uploaded PDF into `<output_folder>/pages_pdf/` and writes per-page
    text files into `<output_folder>/pages_pdf_txt/`. Tries `pdfplumber` first,
    then falls back to `PyPDF2`.
    """
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Uploaded file is not a PDF")

    pathlib.Path(output_folder).mkdir(parents=True, exist_ok=True)
    pages_pdf_dir = os.path.join(output_folder, 'pages_pdf')
    pathlib.Path(pages_pdf_dir).mkdir(parents=True, exist_ok=True)

    save_path = os.path.join(pages_pdf_dir, file.filename)
    contents = await file.read()
    try:
        with open(save_path, 'wb') as wf:
            wf.write(contents)
    finally:
        await file.close()

    text_pages = []
    # Try pdfplumber first for better layout handling
    try:
        import pdfplumber
        with pdfplumber.open(save_path) as pdf:
            for p in pdf.pages:
                text_pages.append(p.extract_text() or '')
    except Exception:
        # Fallback to PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(save_path)
            for p in reader.pages:
                try:
                    text_pages.append(p.extract_text() or '')
                except Exception:
                    text_pages.append('')
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF extraction failed: {e}")

    pages_txt_dir = os.path.join(output_folder, 'pages_pdf_txt')
    pathlib.Path(pages_txt_dir).mkdir(parents=True, exist_ok=True)

    saved_txt_files = []
    for idx, txt in enumerate(text_pages):
        base = os.path.splitext(file.filename)[0]
        fname = f"{base}_page_{idx+1}.txt"
        path = os.path.join(pages_txt_dir, fname)
        try:
            with open(path, 'w', encoding='utf-8') as tf:
                tf.write(txt)
            saved_txt_files.append(path)
        except Exception:
            # continue on write errors
            pass

    combined_path = os.path.join(pages_txt_dir, os.path.splitext(file.filename)[0] + '_full.txt')
    try:
        with open(combined_path, 'w', encoding='utf-8') as cf:
            cf.write('\n\n'.join(text_pages))
    except Exception:
        pass

    return {
        'message': 'PDF extracted',
        'pdf_path': save_path,
        'pages_txt': saved_txt_files,
        'combined_txt': combined_path
    }


@app.post("/extract-pdf-from-path")
async def extract_pdf_from_path(pdf_path: str, output_folder: str = "downloaded_site"):
    """Extract text from an existing PDF on disk. `pdf_path` may be absolute or
    relative; if relative we'll try `output_folder/pages_pdf/<pdf_path>` as a fallback.
    """
    candidate = pdf_path
    if not os.path.isabs(candidate):
        # Try as given relative to cwd
        candidate = os.path.abspath(candidate)

    if not os.path.isfile(candidate):
        # Try under output_folder/pages_pdf
        alt = os.path.join(output_folder, 'pages_pdf', pdf_path)
        if os.path.isfile(alt):
            candidate = alt
        else:
            raise HTTPException(status_code=400, detail=f"PDF not found: {pdf_path}")

    # Use the same extraction logic as the upload endpoint
    text_pages = []
    try:
        import pdfplumber
        with pdfplumber.open(candidate) as pdf:
            for p in pdf.pages:
                text_pages.append(p.extract_text() or '')
    except Exception:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(candidate)
            for p in reader.pages:
                try:
                    text_pages.append(p.extract_text() or '')
                except Exception:
                    text_pages.append('')
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF extraction failed: {e}")

    pages_txt_dir = os.path.join(output_folder, 'pages_pdf_txt')
    pathlib.Path(pages_txt_dir).mkdir(parents=True, exist_ok=True)

    base = os.path.splitext(os.path.basename(candidate))[0]
    saved_txt_files = []
    for idx, txt in enumerate(text_pages):
        fname = f"{base}_page_{idx+1}.txt"
        path = os.path.join(pages_txt_dir, fname)
        try:
            with open(path, 'w', encoding='utf-8') as tf:
                tf.write(txt)
            saved_txt_files.append(path)
        except Exception:
            pass

    combined_path = os.path.join(pages_txt_dir, base + '_full.txt')
    try:
        with open(combined_path, 'w', encoding='utf-8') as cf:
            cf.write('\n\n'.join(text_pages))
    except Exception:
        pass

    return {
        'message': 'PDF extracted from path',
        'pdf_path': candidate,
        'pages_txt': saved_txt_files,
        'combined_txt': combined_path
    }


@app.post("/scrape-process")
async def scrape_process(req: ScrapeProcessRequest):
    """Process previously downloaded per-page .txt files and generate per-page Markdown using the selected model."""
    out_dir = req.output_folder
    pages_txt_dir = os.path.join(out_dir, "pages_txt")
    if not os.path.isdir(pages_txt_dir):
        raise HTTPException(status_code=400, detail=f"Pages text directory not found: {pages_txt_dir}")

    # Validate model
    model_key = req.model_key
    if model_key not in MODEL_CONFIGS:
        raise HTTPException(status_code=400, detail=f"Model {model_key} not available")

    # Ensure model is loaded
    try:
        load_model_on_demand(model_key)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to load model {model_key}: {str(e)}")

    generated_files = []
    errors = []
    md_output_dir = os.path.join(out_dir, "pages_md")
    pathlib.Path(md_output_dir).mkdir(parents=True, exist_ok=True)

    for fname in sorted(os.listdir(pages_txt_dir)):
        if not fname.endswith('.txt'):
            continue
        txt_path = os.path.join(pages_txt_dir, fname)
        try:
            with open(txt_path, 'r', encoding='utf-8') as rf:
                page_text = rf.read()

            if not page_text.strip():
                continue

            if len(page_text) > req.max_chars:
                page_text = page_text[:req.max_chars]

            # Use user-provided prompt template if present, otherwise use default
            if req.prompt_template:
                template = req.prompt_template
                try:
                    if '{content}' in template:
                        prompt = template.format(content=page_text)
                    else:
                        prompt = template + "\n\nPage content:\n" + page_text
                except Exception:
                    # Fallback to appending content if template formatting fails
                    prompt = (template if template else "") + "\n\nPage content:\n" + page_text
            else:
                page_instruction = (
                    "You are a documentation assistant. Convert the following page content into a compact Markdown document."
                    " Produce a suitable title, a short summary (2-4 sentences), headings, preserve and format code examples as fenced code blocks, and keep links inline.\n\nPage content:\n"
                )
                prompt = page_instruction + page_text
            page_req = PromptRequest(prompt=prompt, max_tokens=req.max_new_tokens)
            out = generate_with_model(page_req, model_key, MODEL_CONFIGS[model_key]['name'])

            base = fname[:-4]
            md_path = os.path.join(md_output_dir, f"{base}.md")
            with open(md_path, 'w', encoding='utf-8') as mf:
                mf.write(out.generated_text)
            generated_files.append(md_path)
        except Exception as e:
            errors.append({'file': txt_path, 'error': str(e)})

    if not generated_files:
        raise HTTPException(status_code=500, detail={"message": "No markdown files were generated", "errors": errors})

    preview = ''
    try:
        with open(generated_files[0], 'r', encoding='utf-8') as pf:
            preview = pf.read()[:2000]
    except Exception:
        preview = ''

    return {
        "message": "Processing complete",
        "output_folder": out_dir,
        "markdown_files": generated_files,
        "model_used": MODEL_CONFIGS[model_key]['name'],
        "preview": preview,
        "errors": errors
    }


class AggregateRequest(BaseModel):
    output_folder: str = "downloaded_site"
    model_key: str = "deepseek_tiny"
    prompt: str = "Provide a concise summary of the following content:\n\n{content}"
    max_chars_per_file: int = 20000
    max_new_tokens_per_file: int = 256
    max_chars_aggregate: int = 100000
    max_new_tokens_aggregate: int = 512
    summary_filename: str = "executive_summary.md"


@app.post("/aggregate-md")
async def aggregate_md(req: AggregateRequest):
    """Process all markdown files under `<output_folder>/pages_md`, extract information per-file using the selected model,
    log token consumption and tokens-per-second to the console, and produce an overall summary based on `prompt`.
    """
    out_dir = req.output_folder
    md_dir = os.path.join(out_dir, 'pages_md')

    if not os.path.isdir(md_dir):
        #create dir 
        #os.makedirs(md_dir, exist_ok=True)
        raise HTTPException(status_code=400, detail=f"Markdown directory not found: {md_dir}")

        #raise HTTPException(status_code=400, detail=f"Markdown directory not found: {md_dir}")

    # Validate model
    model_key = req.model_key
    if model_key not in MODEL_CONFIGS:
        raise HTTPException(status_code=400, detail=f"Model {model_key} not available")

    # Ensure model is loaded
    try:
        load_model_on_demand(model_key)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to load model {model_key}: {str(e)}")

    model = models[model_key]
    tokenizer = tokenizers[model_key]
    device = None
    try:
        first_param = next(model.parameters())
        device = first_param.device
    except Exception:
        device = torch.device('cpu')

    has_chat_template = hasattr(tokenizer, 'chat_template') and tokenizer.chat_template is not None

    import time

    per_file_stats = []
    per_file_summaries = []

    for fname in sorted(os.listdir(md_dir)):
        if not fname.endswith('.md'):
            continue
        path = os.path.join(md_dir, fname)
        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()

            if not content.strip():
                continue

            if len(content) > req.max_chars_per_file:
                content = content[:req.max_chars_per_file]

            # Build prompt for this file
            template = req.prompt
            try:
                if '{content}' in template:
                    prompt_text = template.format(content=content)
                else:
                    prompt_text = template + "\n\nContent:\n" + content
            except Exception:
                prompt_text = template + "\n\nContent:\n" + content

            # Prepare model inputs and measure tokens
            if has_chat_template:
                messages = [{"role": "user", "content": prompt_text}]
                input_text = tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
            else:
                input_text = prompt_text

            model_inputs = tokenizer([input_text], return_tensors='pt')
            input_ids = model_inputs.input_ids.to(device)
            attention_mask = model_inputs.attention_mask.to(device) if 'attention_mask' in model_inputs else None
            input_token_count = input_ids.shape[-1]

            # Generate and time
            start = time.perf_counter()
            with torch.no_grad():
                gen_kwargs = {
                    **({} if attention_mask is None else {}),
                    'max_new_tokens': req.max_new_tokens_per_file,
                    'do_sample': True,
                    'temperature': 0.7,
                    'top_p': 0.8,
                    'pad_token_id': getattr(tokenizer, 'eos_token_id', None) or getattr(tokenizer, 'pad_token_id', None)
                }
                generated_ids = model.generate(input_ids=input_ids, **{k:v for k,v in gen_kwargs.items() if v is not None})
            end = time.perf_counter()

            # Compute output token count
            output_ids = generated_ids[0][len(input_ids[0]):].tolist() if generated_ids is not None else []
            output_token_count = len(output_ids)
            total_tokens = input_token_count + output_token_count
            elapsed = max(1e-6, end - start)
            tps = total_tokens / elapsed

            # Decode generated text
            gen_text = tokenizer.decode(output_ids, skip_special_tokens=True)

            # Log to console
            print(f"File: {fname} | tokens_consumed={total_tokens} | tokens/sec={tps:.2f} | elapsed={elapsed:.3f}s")

            per_file_stats.append({
                'file': fname,
                'tokens_consumed': total_tokens,
                'input_tokens': input_token_count,
                'output_tokens': output_token_count,
                'elapsed_seconds': elapsed,
                'tokens_per_second': tps
            })

            per_file_summaries.append(f"--- {fname} ---\n{gen_text}")

        except Exception as e:
            per_file_stats.append({'file': fname, 'error': str(e)})

    if not per_file_summaries:
        raise HTTPException(status_code=500, detail="No markdown summaries were produced")

    # Aggregate summaries into a single final summary
    combined = "\n\n".join(per_file_summaries)
    if len(combined) > req.max_chars_aggregate:
        combined = combined[:req.max_chars_aggregate]

    # Build final prompt
    final_template = req.prompt
    try:
        if '{content}' in final_template:
            final_prompt_text = final_template.format(content=combined)
        else:
            final_prompt_text = final_template + "\n\nCombined content:\n" + combined
    except Exception:
        final_prompt_text = final_template + "\n\nCombined content:\n" + combined

    # Tokenize final prompt
    if has_chat_template:
        messages = [{"role": "user", "content": final_prompt_text}]
        final_input_text = tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
    else:
        final_input_text = final_prompt_text

    final_inputs = tokenizer([final_input_text], return_tensors='pt')
    final_input_ids = final_inputs.input_ids.to(device)

    start_all = time.perf_counter()
    with torch.no_grad():
        final_generated = model.generate(input_ids=final_input_ids, max_new_tokens=req.max_new_tokens_aggregate, do_sample=True, temperature=0.7, top_p=0.8, pad_token_id=getattr(tokenizer, 'eos_token_id', None) or getattr(tokenizer, 'pad_token_id', None))
    end_all = time.perf_counter()

    final_output_ids = final_generated[0][len(final_input_ids[0]):].tolist()
    final_text = tokenizer.decode(final_output_ids, skip_special_tokens=True)
    final_total_tokens = final_input_ids.shape[-1] + len(final_output_ids)
    final_elapsed = max(1e-6, end_all - start_all)
    final_tps = final_total_tokens / final_elapsed

    print(f"FINAL SUMMARY | tokens_consumed={final_total_tokens} | tokens/sec={final_tps:.2f} | elapsed={final_elapsed:.3f}s")

    # Ensure summary folder exists and write the executive summary file
    summary_dir = os.path.join(md_dir, 'summary')
    pathlib.Path(summary_dir).mkdir(parents=True, exist_ok=True)
    fname = req.summary_filename or "executive_summary.md"
    if not fname.lower().endswith('.md'):
        fname = fname + '.md'
    summary_path = os.path.join(summary_dir, fname)

    try:
        from datetime import datetime
        header = f"# Executive Summary\n\nGenerated: {datetime.utcnow().isoformat()} UTC\n\nPrompt: {req.prompt}\n\n"
        stats_section = "## Per-file stats\n\n"
        for s in per_file_stats:
            if 'error' in s:
                stats_section += f"- {s.get('file')}: ERROR - {s.get('error')}\n"
            else:
                stats_section += f"- {s.get('file')}: tokens={s.get('tokens_consumed')} (in={s.get('input_tokens')}, out={s.get('output_tokens')}), elapsed={s.get('elapsed_seconds'):.3f}s, tps={s.get('tokens_per_second'):.2f}\n"

        final_section = "## Final Summary\n\n" + final_text + "\n"

        with open(summary_path, 'w', encoding='utf-8') as sf:
            sf.write(header)
            sf.write(stats_section + "\n")
            sf.write(final_section)

    except Exception as e:
        # If writing fails, return the error but still include generated text
        return {
            'message': 'Aggregation completed but failed to write summary file',
            'error': str(e),
            'per_file_stats': per_file_stats,
            'final_summary': final_text,
            'final_stats': {
                'tokens_consumed': final_total_tokens,
                'elapsed_seconds': final_elapsed,
                'tokens_per_second': final_tps
            }
        }

    return {
        'message': 'Aggregation complete',
        'per_file_stats': per_file_stats,
        'final_summary': final_text,
        'final_stats': {
            'tokens_consumed': final_total_tokens,
            'elapsed_seconds': final_elapsed,
            'tokens_per_second': final_tps
        },
        'summary_file': summary_path
    }


class MergeMdToDocxRequest(BaseModel):
    output_folder: str = "downloaded_site"
    docx_filename: str = "merged.docx"


@app.post("/pages-md/to-docx")
async def merge_pages_md_to_docx(req: MergeMdToDocxRequest):
    """Merge all `.md` files under `<output_folder>/pages_md` into a single Word `.docx` file.

    The endpoint writes the output to `<output_folder>/pages_md/merged/<docx_filename>` and returns that path.
    Requires `python-docx` (`pip install python-docx`).
    """
    out_dir = req.output_folder
    md_dir = os.path.join(out_dir, 'pages_md')
    if not os.path.isdir(md_dir):
        raise HTTPException(status_code=400, detail=f"Markdown directory not found: {md_dir}")

    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as e:
        raise HTTPException(status_code=500, detail=("python-docx is required for this endpoint. Install with: pip install python-docx\n" + str(e)))

    merged_dir = os.path.join(md_dir, 'merged')
    pathlib.Path(merged_dir).mkdir(parents=True, exist_ok=True)

    fname = req.docx_filename or 'merged.docx'
    if not fname.lower().endswith('.docx'):
        fname = fname + '.docx'

    out_path = os.path.join(merged_dir, fname)

    doc = Document()

    files = sorted([f for f in os.listdir(md_dir) if f.endswith('.md')])
    if not files:
        raise HTTPException(status_code=400, detail=f"No markdown (.md) files found in {md_dir}")

    first = True
    for md_fname in files:
        md_path = os.path.join(md_dir, md_fname)
        try:
            with open(md_path, 'r', encoding='utf-8') as rf:
                lines = rf.readlines()
        except Exception as e:
            # Skip unreadable files but log
            print(f"Failed to read {md_path}: {e}")
            continue

        if not first:
            # insert a page break between files
            try:
                doc.add_page_break()
            except Exception:
                pass
        first = False

        # Add a heading for the file (filename-based)
        title = os.path.splitext(md_fname)[0].replace('_', ' ')
        doc.add_heading(title, level=1)

        in_code = False
        for raw in lines:
            line = raw.rstrip('\n')
            if line.strip().startswith('```'):
                in_code = not in_code
                continue
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run(line)
                try:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                except Exception:
                    pass
                continue

            # Headings
            if line.startswith('#'):
                # count leading hashes
                hashes = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                level = min(max(hashes, 1), 4)
                doc.add_heading(text, level=level)
                continue

            # Horizontal rule
            if line.strip().startswith('---'):
                doc.add_paragraph('')
                continue

            # Normal paragraph
            if line.strip() == '':
                # preserve paragraph breaks
                doc.add_paragraph('')
            else:
                doc.add_paragraph(line)

    try:
        doc.save(out_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to write docx: {e}")

    return {
        'message': 'Merged markdown files into docx',
        'output_docx': out_path,
        'num_files': len(files)
    }

@app.post("/generate/qwen", response_model=TextResponse)
async def generate_text_qwen(request: PromptRequest):
    """
    Generate text using Qwen2.5-Coder 1.5B model (distilled version)
    """
    return generate_with_model(request, "qwen", "Qwen2.5-Coder-1.5B-Instruct")

@app.post("/generate/deepseek", response_model=TextResponse)
async def generate_text_deepseek(request: PromptRequest):
    """
    Generate text using DeepSeek-Coder-V2-Lite Base model
    """
    return generate_with_model(request, "deepseek", "DeepSeek-Coder-V2-Lite-Base")

@app.post("/generate/qwen-tiny", response_model=TextResponse)
async def generate_text_qwen_tiny(request: PromptRequest):
    """
    Generate text using Qwen2.5-Coder 0.5B model (ultra-lightweight)
    """
    return generate_with_model(request, "qwen_tiny", "Qwen2.5-Coder-0.5B-Instruct")

@app.post("/generate/codellama", response_model=TextResponse)
async def generate_text_codellama(request: PromptRequest):
    """
    Generate text using Meta's CodeLlama 7B Python model
    """
    return generate_with_model(request, "codellama", "CodeLlama-7B-Python")

@app.post("/generate/deepseek-tiny", response_model=TextResponse)
async def generate_text_deepseek_tiny(request: PromptRequest):
    """
    Generate text using DeepSeek-Coder 1.3B model (smallest version)
    """
    return generate_with_model(request, "deepseek_tiny", "DeepSeek-Coder-1.3B-Base")

@app.post("/generate", response_model=TextResponse)
async def generate_text_default(request: PromptRequest):
    """
    Generate text using default Qwen 1.5B model (for backward compatibility)
    """
    return generate_with_model(request, "qwen", "Qwen2.5-Coder-1.5B-Instruct")

@app.get("/")
async def root():
    """Health check endpoint"""
    return {
        "message": "Multi-Model Coder API with Distilled Models",
        "version": "2.0 - Distilled & Optimized",
        "available_endpoints": [
            "/generate/qwen - Qwen2.5-Coder-1.5B (Fast & Efficient)",
            "/generate/qwen-tiny - Qwen2.5-Coder-0.5B (Ultra-lightweight)",
            "/generate/deepseek - DeepSeek-Coder-V2-Lite-Base",
            "/generate/deepseek-tiny - DeepSeek-Coder-1.3B (Smallest DeepSeek)",
            "/generate/codellama - Meta CodeLlama-7B-Python (Open Source)",
            "/generate - Default (Qwen 1.5B)"
        ],
        "benefits": [
            "🚀 Faster inference with distilled models",
            "💾 Lower memory usage (~1-3GB vs 6-8GB)",
            "⚡ Quicker model loading times",
            "🎯 Optimized for code generation"
        ]
    }

@app.get("/health")
async def health_check():
    """Check if models are loaded and ready with detailed GPU information"""
    gpu_info = get_gpu_info()
    
    return {
        "status": "ready",
        "loading_strategy": "lazy_loading",
        "models": {
            model_key: {
                "status": loading_status.get(model_key, "not_loaded"),
                "in_memory": model_key in models,
                "device_info": model_device_info.get(model_key, {"status": "not_loaded"})
            }
            for model_key in MODEL_CONFIGS.keys()
        },
        "gpu_info": gpu_info,
        "system_memory": {
            "total_gb": f"{psutil.virtual_memory().total / 1024**3:.2f}",
            "available_gb": f"{psutil.virtual_memory().available / 1024**3:.2f}",
            "used_percent": f"{psutil.virtual_memory().percent:.1f}"
        }
    }

@app.get("/models")
async def list_models():
    """List available models and their status"""
    return {
        "loading_strategy": "lazy_loading - models load on first request",
        "memory_optimization": "only one model in memory at a time",
        "available_models": {
            model_key: {
                "name": config["name"],
                "endpoint": f"/generate/{model_key}",
                "status": loading_status.get(model_key, "not_loaded"),
                "in_memory": model_key in models,
                "config": {k: str(v) for k, v in config["config"].items()}
            }
            for model_key, config in MODEL_CONFIGS.items()
        }
    }

@app.post("/models/{model_key}/preload")
async def preload_model(model_key: str):
    """Preload a specific model into memory"""
    try:
        load_model_on_demand(model_key)
        return {
            "message": f"Model {model_key} preloaded successfully",
            "model": model_key,
            "status": loading_status.get(model_key)
        }
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to preload model {model_key}: {str(e)}")

@app.post("/models/{model_key}/unload")
async def unload_model(model_key: str):
    """Unload a specific model from memory"""
    if model_key in models:
        # Capture memory before unloading
        memory_before = get_gpu_info() if torch.cuda.is_available() else None
        
        del models[model_key]
        del tokenizers[model_key]
        if model_key in model_device_info:
            del model_device_info[model_key]
        loading_status[model_key] = "not_loaded"
        
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
        
        # Capture memory after unloading
        memory_after = get_gpu_info() if torch.cuda.is_available() else None
        
        return {
            "message": f"Model {model_key} unloaded successfully",
            "model": model_key,
            "gpu_memory_freed": True if torch.cuda.is_available() else False,
            "memory_before": memory_before,
            "memory_after": memory_after
        }
    else:
        return {
            "message": f"Model {model_key} was not loaded",
            "model": model_key
        }


class FineTuneRequest(BaseModel):
    # Provide either `training_examples` (list of {prompt, completion}) or `training_file` (local JSONL path)
    training_examples: list = None
    training_file: str = None
    use_lora: bool = True
    lora_r: int = 8
    lora_alpha: int = 32
    lora_dropout: float = 0.05
    num_train_epochs: int = 1
    per_device_train_batch_size: int = 4
    output_dir: str = None  # if not set, will use fine_tuned/<model_key>
    load_in_8bit: bool = True
    gradient_checkpointing: bool = True
    max_seq_length: int = 1024


@app.post("/models/{model_key}/fine-tune")
async def fine_tune_model(model_key: str, req: FineTuneRequest):
    """Enqueue a fine-tune job and run it in background.

    Returns a job id that can be polled via `/jobs/{job_id}`.
    """
    if model_key not in MODEL_CONFIGS:
        raise HTTPException(status_code=400, detail=f"Model {model_key} not available")

    # Create job entry
    job_id = str(uuid.uuid4())
    jobs[job_id] = {
        'id': job_id,
        'model_key': model_key,
        'status': 'queued',
        'created_at': time.time(),
        'started_at': None,
        'finished_at': None,
        'logs': [],
        'result': None
    }

    _log_job(job_id, f"Job queued for model {model_key}")

    # Schedule background execution
    loop = asyncio.get_event_loop()

    async def _run_job():
        jobs[job_id]['status'] = 'running'
        jobs[job_id]['started_at'] = time.time()
        _log_job(job_id, 'Starting fine-tune job')

        # Run the existing synchronous training logic in a thread pool
        try:
            from concurrent.futures import ThreadPoolExecutor

            def worker():
                try:
                    return _fine_tune_run(model_key, req, job_id)
                except Exception as e:
                    return {'error': str(e)}

            with ThreadPoolExecutor(max_workers=1) as ex:
                fut = ex.submit(worker)
                res = fut.result()

            if isinstance(res, dict) and res.get('error'):
                _log_job(job_id, 'Job failed: ' + res.get('error'))
                jobs[job_id]['status'] = 'failed'
                jobs[job_id]['result'] = {'error': res.get('error')}
            else:
                _log_job(job_id, 'Job completed successfully')
                jobs[job_id]['status'] = 'finished'
                jobs[job_id]['result'] = res

        except Exception as e:
            _log_job(job_id, 'Unhandled exception: ' + str(e))
            jobs[job_id]['status'] = 'failed'
            jobs[job_id]['result'] = {'error': str(e)}
        finally:
            jobs[job_id]['finished_at'] = time.time()

    loop.create_task(_run_job())

    return {'job_id': job_id, 'status': 'queued'}


@app.get("/jobs")
async def list_jobs():
    """List all background jobs with summary information"""
    return {
        'jobs': [
            {
                'id': j['id'],
                'model_key': j['model_key'],
                'status': j['status'],
                'created_at': j['created_at'],
                'started_at': j.get('started_at'),
                'finished_at': j.get('finished_at')
            }
            for j in jobs.values()
        ]
    }


@app.get("/jobs/{job_id}")
async def get_job(job_id: str):
    entry = jobs.get(job_id)
    if not entry:
        raise HTTPException(status_code=404, detail=f"Job {job_id} not found")
    return entry


class QueryFineTunedRequest(BaseModel):
    fine_model_path: str
    prompt: str
    max_new_tokens: int = 256
    temperature: float = 0.7
    top_p: float = 0.8
    unload_after: bool = True


@app.post("/models/query-finetuned")
async def query_finetuned(req: QueryFineTunedRequest):
    """Load a fine-tuned model from `fine_model_path` (local dir) and run a prompt against it.

    The model/tokenizer will be loaded into memory (cached) under a derived key. If `unload_after` is true,
    the model will be removed from memory after generation to free resources.
    """
    path = req.fine_model_path
    if not path:
        raise HTTPException(status_code=400, detail="fine_model_path is required")

    # Derive a safe cache key for this path
    key = f"finetuned::{os.path.abspath(path)}"

    # Load tokenizer/model if not already loaded
    if key not in tokenizers or key not in models:
        try:
            tokenizers[key] = AutoTokenizer.from_pretrained(path, trust_remote_code=False)
            models[key] = AutoModelForCausalLM.from_pretrained(
                path,
                device_map='auto',
                torch_dtype=(torch.float16 if torch.cuda.is_available() else None),
                low_cpu_mem_usage=True
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to load fine-tuned model from {path}: {e}")

    model = models[key]
    tokenizer = tokenizers[key]

    # Prepare prompt
    try:
        model_inputs = tokenizer([req.prompt], return_tensors="pt")
        # Move inputs to model device
        try:
            first = next(model.parameters())
            device = first.device
        except Exception:
            device = torch.device('cpu')
        model_inputs = {k: v.to(device) for k, v in model_inputs.items()}

        with torch.no_grad():
            generated = model.generate(
                **model_inputs,
                max_new_tokens=req.max_new_tokens,
                do_sample=True,
                temperature=req.temperature,
                top_p=req.top_p,
                pad_token_id=getattr(tokenizer, 'eos_token_id', None) or getattr(tokenizer, 'pad_token_id', None)
            )

        input_len = model_inputs['input_ids'].shape[-1]
        output_ids = generated[0][input_len:].tolist()
        generated_text = tokenizer.decode(output_ids, skip_special_tokens=True)

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation failed: {e}")

    # Optionally unload
    if req.unload_after:
        try:
            del models[key]
            del tokenizers[key]
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
        except Exception:
            pass

    return {
        'generated_text': generated_text,
        'model_path': path
    }

@app.get("/gpu")
async def gpu_status():
    """Get detailed GPU status and model locations"""
    gpu_info = get_gpu_info()
    
    model_locations = {}
    for model_key in models.keys():
        device_info = detect_model_device(models[model_key])
        model_locations[model_key] = {
            "model_name": MODEL_CONFIGS[model_key]["name"],
            "device_info": device_info,
            "status": loading_status.get(model_key, "unknown")
        }
    
    return {
        "gpu_hardware": gpu_info,
        "loaded_models": model_locations,
        "recommendations": {
            "models_on_gpu": len([m for m in model_locations.values() if "cuda" in str(m["device_info"].get("primary_device", ""))]),
            "total_models_loaded": len(model_locations),
            "gpu_memory_efficient": all(gpu["utilization_percent"] != "0.0" for gpu in gpu_info.get("devices", []) if gpu["utilization_percent"] != "0.0") if gpu_info.get("cuda_available") else False
        }
    }

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)